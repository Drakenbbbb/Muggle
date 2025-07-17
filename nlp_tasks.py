# ==============================================================================
# 短视频内容标签与热点分析系统 (V3.0 高级版)
#
# 核心功能:
# 1. 热点趋势总览：多维度分析热门视频、标签与创作者。
# 2. 爆款密码挖掘：使用Apriori算法挖掘高收益的标签组合。
# 3. 内容策略分析：对头部创作者进行雷达图对比分析。
# 4. 智能报告生成：一键导出所有分析结果为Word文档。
# ==============================================================================

import os
import streamlit as st
import pandas as pd
import numpy as np
import ast
import re
import jieba
import lightgbm as lgb
import plotly.express as px
import plotly.graph_objects as go
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.model_selection import train_test_split
from sklearn.metrics import classification_report, accuracy_score, f1_score, roc_auc_score
from mlxtend.preprocessing import TransactionEncoder
from mlxtend.frequent_patterns import apriori, association_rules
from docx import Document
from docx.shared import Inches
import io
import matplotlib.pyplot as plt
import seaborn as sns
from pathlib import Path
import matplotlib
import platform

# 设置matplotlib中文字体，自动适配操作系统
def set_chinese_font():
    if platform.system() == 'Windows':
        matplotlib.rcParams['font.sans-serif'] = ['SimHei']  # 黑体
    elif platform.system() == 'Darwin':
        matplotlib.rcParams['font.sans-serif'] = ['Heiti TC']  # macOS
    else:
        matplotlib.rcParams['font.sans-serif'] = ['SimHei', 'Arial Unicode MS']  # Linux常用
    matplotlib.rcParams['axes.unicode_minus'] = False  # 负号正常显示

# ==============================================================================
# 0. 页面配置 (必须是第一个st命令)
# ==============================================================================
st.set_page_config(page_title="B站内容决策工具", layout="wide", initial_sidebar_state="expanded")


# ==============================================================================
# 1. 数据加载与预处理模块 (唯一正确版本)
# ==============================================================================
@st.cache_data
def load_data(csv_path):
    """
    加载并预处理B站热门视频数据。
    这是系统中唯一的数据加载函数，确保所有模块使用统一的数据源。
    """
    try:
        df = pd.read_csv(csv_path)

        # 核心修正：将upload_time转换为created_at，并处理潜在错误
        if 'upload_time' in df.columns:
            # errors='coerce' 会将无法解析的日期变为NaT（Not a Time），而不是报错
            df['created_at'] = pd.to_datetime(df['upload_time'], errors='coerce')
            
            # 检查是否有转换失败的行
            if df['created_at'].isnull().any():
                st.warning("部分行的 'upload_time' 格式不正确，已跳过。")
                df.dropna(subset=['created_at'], inplace=True) # 移除转换失败的行
        else:
            # 如果完全没有该列，则无法进行时间分析
            st.error("错误：数据文件中缺少 'upload_time' 列，无法进行趋势分析。")
            return pd.DataFrame()

        # 安全地将tags字符串解析为list
        if 'tags' in df.columns:
            df['tags_list'] = df['tags'].apply(lambda x: ast.literal_eval(x) if isinstance(x, str) and x.startswith('[') else [])
        else:
            df['tags_list'] = [[] for _ in range(len(df))]

        # 清洗和转换数据类型
        df['title'] = df['title'].apply(lambda x: str(x) if pd.notnull(x) else '')
        df['author'] = df['author'].astype(str)

        # 计算互动量 (点赞+评论+投币)
        for col in ['likes', 'comments', 'coins']:
            if col not in df.columns:
                df[col] = 0
        df['engagement'] = df['likes'].astype(int) + df['comments'].astype(int) + df['coins'].astype(int)

        # 计算互动率 (互动量/播放量)，处理分母为0的情况
        if 'views' not in df.columns:
            df['views'] = 0
        df['engagement_rate'] = (df['engagement'] / df['views'].replace(0, np.nan)).fillna(0)

        # 创建主标签
        df['main_tag'] = df['tags_list'].apply(lambda x: x[0] if x else '未知')

        # 创建用于NLP的组合文本特征
        df['text_for_nlp'] = df['title'] + ' ' + df['tags_list'].apply(lambda x: ' '.join(x))

        return df

    except FileNotFoundError:
        st.error(f"错误：数据文件未找到，请确认路径 '{csv_path}' 是否正确。")
        return pd.DataFrame()
    except Exception as e:
        st.error(f"加载或处理数据时发生未知错误：{e}")
        return pd.DataFrame()


# ==============================================================================
# 2. 核心分析与建模模块 (带缓存)
# ==============================================================================
@st.cache_resource  # 缓存训练好的模型
def get_text_classifier(df, min_count=5):
    """训练并缓存一个文本分类器。"""
    tag_counts = df['main_tag'].value_counts()
    valid_tags = tag_counts[tag_counts >= min_count].index
    df_filtered = df[df['main_tag'].isin(valid_tags)].copy()
    X = df_filtered['text_for_nlp']
    y = df_filtered['main_tag']
    vectorizer = TfidfVectorizer(tokenizer=lambda x: jieba.cut(x), max_features=2000)
    X_vec = vectorizer.fit_transform(X)
    X_train, X_test, y_train, y_test = train_test_split(X_vec, y, test_size=0.25, random_state=42, stratify=y)
    clf = lgb.LGBMClassifier(random_state=42)
    clf.fit(X_train, y_train)
    return clf, vectorizer, X_test, y_test, valid_tags


@st.cache_resource
def get_hot_predictor(df, top_n=100):
    """训练并缓存一个爆款潜力预测器。"""
    df_copy = df.copy()
    df_copy['is_hot'] = 0
    df_copy.iloc[:top_n, df_copy.columns.get_loc('is_hot')] = 1
    X = df_copy['text_for_nlp']
    y = df_copy['is_hot']
    vectorizer = TfidfVectorizer(tokenizer=lambda x: jieba.cut(x), max_features=2000)
    X_vec = vectorizer.fit_transform(X)
    X_train, X_test, y_train, y_test = train_test_split(X_vec, y, test_size=0.25, random_state=42, stratify=y)
    clf = lgb.LGBMClassifier(random_state=42)
    clf.fit(X_train, y_train)
    return clf, vectorizer, X_test, y_test


def analyze_top_authors(df, top_n=10):
    """分析头部创作者的核心指标。"""
    author_stats = df.groupby('author').agg(
        video_count=('title', 'count'),
        total_views=('views', 'sum'),
        total_likes=('likes', 'sum'),
        total_engagement=('engagement', 'sum'),
        avg_engagement_rate=('engagement_rate', 'mean')
    ).sort_values('total_views', ascending=False).head(top_n)
    return author_stats


@st.cache_data
def get_association_rules(df, min_support=0.01, min_confidence=0.1):
    """使用Apriori算法挖掘标签关联规则。"""
    # 过滤掉没有标签的视频
    tags_df = df[df['tags_list'].apply(len) > 0]

    # 数据转换
    te = TransactionEncoder()
    te_ary = te.fit(tags_df['tags_list']).transform(tags_df['tags_list'])
    rules_df = pd.DataFrame(te_ary, columns=te.columns_)

    # 挖掘频繁项集
    frequent_itemsets = apriori(rules_df, min_support=min_support, use_colnames=True)
    if frequent_itemsets.empty:
        return pd.DataFrame()  # 如果没有找到频繁项集，返回空表

    # 计算关联规则
    rules = association_rules(frequent_itemsets, metric="confidence", min_threshold=min_confidence)
    rules = rules.sort_values(['lift', 'confidence'], ascending=[False, False])
    return rules


# ==============================================================================
# 3. 报告生成模块
# ==============================================================================
def generate_word_report(df, top_authors_stats, top_tags_stats, rules):
    """生成并返回一个包含核心分析的Word文档的内存字节流。"""
    document = Document()
    document.add_heading('B站热门内容趋势分析报告', 0)

    # --- 内容创作建议 ---
    document.add_heading('一、核心内容创作建议', level=1)
    document.add_paragraph(generate_insight(df), style='List Bullet')
    document.add_paragraph(f"热门内容主要集中在'{top_tags_stats.index[0]}'、'{top_tags_stats.index[1]}'等领域。",
                           style='List Bullet')
    document.add_paragraph(f"头部创作者如'{top_authors_stats.index[0]}'在播放量上具有显著优势，其内容策略值得借鉴。",
                           style='List Bullet')
    if not rules.empty:
        top_rule = rules.iloc[0]
        antecedent = ', '.join(list(top_rule['antecedents']))
        consequent = ', '.join(list(top_rule['consequents']))
        document.add_paragraph(
            f"数据显示，'{antecedent}'与'{consequent}'的组合具有很高的关联性 (提升度: {top_rule['lift']:.2f})，是潜在的爆款密码。",
            style='List Bullet')

    # --- 热门标签分析 ---
    document.add_heading('二、热门标签分析', level=1)
    p = document.add_paragraph('下图展示了热门视频中出现次数最多的Top 20标签。')
    fig1 = px.bar(top_tags_stats, x=top_tags_stats.index, y='count', title="热门标签Top 20 (按出现次数)",
                  labels={'x': '标签', 'y': '出现次数'})
    img_buffer = io.BytesIO()
    fig1.write_image(img_buffer, format="png")
    img_buffer.seek(0)
    document.add_picture(img_buffer, width=Inches(6.0))

    # --- 头部创作者分析 ---
    document.add_heading('三、头部创作者分析', level=1)
    p = document.add_paragraph('下表展示了总播放量最高的Top 10创作者的核心数据。')
    # 添加表格
    table_df = top_authors_stats.reset_index()
    table = document.add_table(rows=1, cols=len(table_df.columns))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(table_df.columns):
        hdr_cells[i].text = col_name
    for index, row in table_df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)

    # --- 总结 ---
    document.add_heading('四、总结', level=1)
    document.add_paragraph(f"本次分析基于 {len(df)} 条热门视频数据。核心发现包括：")
    document.add_paragraph(f"热门内容主要集中在'{top_tags_stats.index[0]}'、'{top_tags_stats.index[1]}'等领域。",
                           style='List Bullet')
    document.add_paragraph(f"头部创作者如'{top_authors_stats.index[0]}'在播放量上具有显著优势。", style='List Bullet')

    # 将文档保存到内存字节流
    doc_buffer = io.BytesIO()
    document.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer


# ==============================================================================
# 4. Streamlit UI 页面渲染
# ==============================================================================

# --- 页面1: 热点趋势总览 ---
def page_overview(df):
    st.title("📈 热点趋势总览")
    st.markdown("---")

    col1, col2, col3, col4 = st.columns(4)
    col1.metric("视频总数", f"{len(df):,}")
    col2.metric("总播放量", f"{df['views'].sum():,.0f}")
    col3.metric("总点赞量", f"{df['likes'].sum():,.0f}")
    col4.metric("创作者总数", f"{df['author'].nunique():,}")

    st.markdown("### 热门标签分析")
    tag_stats = df.explode('tags_list').groupby('tags_list').agg(
        count=('title', 'count'),
        avg_views=('views', 'mean')
    ).sort_values('count', ascending=False)

    top_tags_by_count = tag_stats['count'].head(20)
    fig1 = px.bar(top_tags_by_count, x=top_tags_by_count.index, y=top_tags_by_count.values,
                  title="热门标签Top 20 (按出现次数)", labels={'x': '标签', 'y': '出现次数'})
    st.plotly_chart(fig1, use_container_width=True)

    st.markdown("### 头部创作者分析 (按总播放量)")
    top_authors_stats = analyze_top_authors(df)
    st.dataframe(top_authors_stats.style.format("{:,.0f}"))


# --- 页面2: 爆款密码挖掘 (Apriori) ---
def page_association_rules(df):
    st.title("🔑 爆款密码挖掘 (Apriori算法)")
    st.info("挖掘能带来流量增益的'标签黄金组合'。Lift(提升度) > 1 表示正向关联。")

    col1, col2 = st.columns([1, 3])
    with col1:
        st.subheader("参数调整")
        min_support = st.slider("最小支持度 (Min Support)", 0.001, 0.1, 0.01, 0.001, format="%.3f")
        min_confidence = st.slider("最小置信度 (Min Confidence)", 0.05, 0.5, 0.1, 0.05)

    rules = get_association_rules(df, min_support, min_confidence)

    with col2:
        st.subheader("关联规则结果")
        if not rules.empty:
            display_rules = rules[['antecedents', 'consequents', 'support', 'confidence', 'lift']].copy()
            display_rules['antecedents'] = display_rules['antecedents'].apply(lambda x: ', '.join(list(x)))
            display_rules['consequents'] = display_rules['consequents'].apply(lambda x: ', '.join(list(x)))
            st.dataframe(
                display_rules.head(20).style.format({'support': '{:.3f}', 'confidence': '{:.3f}', 'lift': '{:.2f}'}))
        else:
            st.warning("在当前参数下未找到关联规则，请尝试降低最小支持度或最小置信度。")


# --- 页面3: 内容策略分析 ---
def page_strategy_analysis(df):
    st.title("🎯 内容策略分析")

    st.subheader("头部创作者内容策略雷达图")
    top_authors_stats = analyze_top_authors(df)

    # 数据归一化，用于雷达图
    scaler = lambda x: (x - x.min()) / (x.max() - x.min()) if x.max() - x.min() != 0 else x / x.max()
    normalized_stats = top_authors_stats.apply(scaler).reset_index()

    fig = go.Figure()
    for index, row in normalized_stats.iterrows():
        fig.add_trace(go.Scatterpolar(
            r=row.drop('author').values,
            theta=row.drop('author').index,
            fill='toself',
            name=row['author']
        ))

    fig.update_layout(
        polar=dict(radialaxis=dict(visible=True, range=[0, 1])),
        showlegend=True,
        title="Top 10创作者核心指标对比"
    )
    st.plotly_chart(fig, use_container_width=True)


# --- 新增页面4: AI模型中心 ---
def page_models_center(df):
    st.title("🤖 AI模型中心")
    st.info("在这里，你可以评估系统内置的AI模型性能，并进行交互式预测。")

    # --- 内容领域分类模型 ---
    with st.container(border=True):
        st.subheader("🏷️ 内容领域分类模型")
        clf, vectorizer, X_test, y_test, valid_tags = get_text_classifier(df)
        y_pred = clf.predict(X_test)

        st.markdown("**模型性能:**")
        col1, col2 = st.columns(2)
        col1.metric("模型准确率 (Accuracy)", f"{accuracy_score(y_test, y_pred):.2%}")
        col2.metric("宏平均 F1-Score", f"{f1_score(y_test, y_pred, average='macro', zero_division=0):.3f}")
        with st.expander("点击查看详细分类报告"):
            st.text(classification_report(y_test, y_pred, digits=3, zero_division=0))

        with st.form(key='classifier_form'):
            user_input = st.text_input('输入视频标题或标签:', '介绍一款最新的AI工具')
            submitted = st.form_submit_button('预测主标签')
            if submitted and user_input:
                pred_tag = clf.predict(vectorizer.transform([user_input]))[0]
                st.success(f'预测主标签: **{pred_tag}**')

    # --- 爆款潜力预测模型 ---
    with st.container(border=True):
        st.subheader("🚀 爆款潜力预测模型")
        hot_clf, hot_vectorizer, X_test_hot, y_test_hot = get_hot_predictor(df)
        y_pred_hot = hot_clf.predict(X_test_hot)
        y_prob_hot = hot_clf.predict_proba(X_test_hot)[:, 1]

        st.markdown("**模型性能:**")
        col1, col2, col3 = st.columns(3)
        col1.metric("模型准确率 (Accuracy)", f"{accuracy_score(y_test_hot, y_pred_hot):.2%}")
        col2.metric("F1-Score", f"{f1_score(y_test_hot, y_pred_hot, zero_division=0):.3f}")
        col3.metric("AUC", f"{roc_auc_score(y_test_hot, y_prob_hot):.3f}")

        with st.form(key='predictor_form'):
            user_input2 = st.text_input('输入你的内容创意:', '挑战24小时只吃便利店食物')
            submitted2 = st.form_submit_button('预测爆款潜力')
            if submitted2 and user_input2:
                prob = hot_clf.predict_proba(hot_vectorizer.transform([user_input2]))[0, 1]
                st.progress(prob)
                st.success(f'**爆款概率: {prob:.2%}**')


# --- 页面5: 报告生成与下载 ---
def page_report_generation(df):
    st.title("📄 智能报告生成")
    st.info("一键生成包含核心图表与数据的分析报告。注意：报告生成（特别是第一次）可能需要10-20秒，请耐心等待。")

    if st.button("生成分析报告 (Word)"):
        with st.spinner("正在分析数据并渲染图表，请稍候..."):
            tag_stats = df.explode('tags_list').groupby('tags_list').agg(count=('title', 'count')).sort_values('count',
                                                                                                               ascending=False)
            top_authors_stats = analyze_top_authors(df)
            rules = get_association_rules(df)

            report_buffer = generate_word_report(df, top_authors_stats, tag_stats.head(20), rules)

            st.success("报告生成完毕！点击下方按钮下载。")
            st.download_button(
                label="下载Word报告",
                data=report_buffer,
                file_name="B站热点分析报告.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )


# --- 趋势预测页面 ---
def page_trend_forecast(df):
    st.title("📊 趋势预测与热点挖掘")
    st.info("基于历史数据，自动识别近期高潜力标签和内容方向，辅助选题决策。")

    if 'created_at' not in df.columns or df['created_at'].isnull().all():
        st.warning("数据中缺少有效的 'created_at' 时间列，无法进行趋势分析。")
        return

    # 1. 标签热度趋势
    st.subheader("标签热度趋势")
    
    # 修正：将 'created_at' 设置为索引以进行时间序列分组
    trend_df = df.set_index('created_at')
    tag_trend = trend_df.explode('tags_list').groupby(['tags_list', pd.Grouper(freq='7D')])['views'].sum().reset_index()
    
    if not tag_trend.empty:
        top_tags = tag_trend.groupby('tags_list')['views'].sum().nlargest(10).index.tolist()
        tag_trend_top = tag_trend[tag_trend['tags_list'].isin(top_tags)]
        fig = px.line(tag_trend_top, x='created_at', y='views', color='tags_list', markers=True, title="近30天热门标签趋势")
        st.plotly_chart(fig, use_container_width=True)
    else:
        st.info("数据量不足，无法生成标签热度趋势图。")

    # 2. 潜力标签推荐（近7天增速最快）
    st.subheader("高潜力标签推荐")
    recent = df[df['created_at'] > df['created_at'].max() - pd.Timedelta(days=7)]
    last = df[(df['created_at'] <= df['created_at'].max() - pd.Timedelta(days=7)) & (df['created_at'] > df['created_at'].max() - pd.Timedelta(days=14))]
    
    if not recent.empty and not last.empty:
        recent_tags = recent.explode('tags_list')['tags_list'].value_counts()
        last_tags = last.explode('tags_list')['tags_list'].value_counts()
        growth = (recent_tags - last_tags).sort_values(ascending=False).dropna().head(10)
        st.dataframe(growth.rename('近7天净增长').to_frame())
    else:
        st.info("数据时间跨度不足14天，无法计算潜力标签。")

    # 3. 爆款预警（播放/互动异常增长）
    st.subheader("爆款内容预警")
    
    # 修正：处理标准差为0的情况，避免 'float' object has no attribute 'replace' 错误
    views_std = df['views'].std()
    engage_std = df['engagement'].std()

    # 如果标准差为0（所有值都一样），Z-score也为0，否则正常计算
    df['views_zscore'] = ((df['views'] - df['views'].mean()) / views_std) if views_std > 0 else 0
    df['engage_zscore'] = ((df['engagement'] - df['engagement'].mean()) / engage_std) if engage_std > 0 else 0
    
    hot_df = df[(df['views_zscore'] > 2) | (df['engage_zscore'] > 2)].sort_values('views', ascending=False).head(10)
    st.dataframe(hot_df[['title', 'author', 'views', 'engagement', 'created_at']])


# --- 标签竞争度与长尾标签 ---
def tag_competition_analysis(df):
    st.title("🏷️ 标签竞争度与长尾标签推荐")
    st.info("分析标签的流量潜力和竞争度，推荐高转化长尾标签。")
    tag_stats = df.explode('tags_list').groupby('tags_list').agg(
        count=('title', 'count'),
        avg_views=('views', 'mean'),
        avg_engage=('engagement', 'mean')
    )
    tag_stats['竞争度'] = tag_stats['count'] / tag_stats['avg_views']
    tag_stats['流量潜力'] = tag_stats['avg_views'] * tag_stats['avg_engage']
    st.dataframe(tag_stats.sort_values('流量潜力', ascending=False).head(20))
    st.subheader("长尾标签推荐")
    long_tail = tag_stats[(tag_stats['count'] < tag_stats['count'].median()) & (tag_stats['流量潜力'] > tag_stats['流量潜力'].median())]
    st.dataframe(long_tail.sort_values('流量潜力', ascending=False).head(10))


# --- 竞品内容日历 ---
def competitor_calendar(df):
    st.title("📅 竞品内容日历")
    st.info("可视化竞品账号的内容发布时间和主题分布，辅助优化自身内容策略。")
    
    # 设置中文字体，防止图像中文乱码
    set_chinese_font()

    if 'author' not in df.columns or df['author'].nunique() == 0:
        st.warning("数据中没有创作者信息，无法进行此分析。")
        return

    # 核心改进：预先筛选出有/无有效发布时间的创作者
    authors_with_time_data = df.dropna(subset=['created_at'])['author'].unique()

    if len(authors_with_time_data) == 0:
        st.warning("所有创作者的数据都缺少有效的发布时间，无法生成任何内容日历。")
        st.info("请检查原始数据文件中的 'upload_time' 列是否存在且格式正确（例如 '2023-10-26 18:00:00'）。")
        return

    all_authors_in_df = df['author'].unique()
    authors_without_time_data = [author for author in all_authors_in_df if author not in authors_with_time_data]

    # 让用户从有数据的作者列表中选择
    author = st.selectbox("选择要分析的竞品账号：", authors_with_time_data)

    # 如果有作者因为数据缺失而未显示，则以可展开的方式告知用户
    if authors_without_time_data:
        with st.expander(f"ℹ️ 注意：有 {len(authors_without_time_data)} 位创作者因缺少发布时间数据未在列表中显示"):
            st.write("以下创作者无法进行发布时间分析，因为在数据源中找不到他们视频的有效发布日期：")
            # 显示部分列表以防过长
            st.json(authors_without_time_data[:20])

    # 使用 .copy() 避免 SettingWithCopyWarning
    author_df = df[df['author'] == author].copy()

    # 此时 author_df 必然有有效的 'created_at'，无需再次检查
    author_df['date'] = author_df['created_at'].dt.date
    author_df['hour'] = author_df['created_at'].dt.hour

    # 使用 pivot_table 创建热力图数据
    pivot = author_df.pivot_table(index='date', columns='hour', values='views', aggfunc='sum', fill_value=0)

    if pivot.empty:
        st.info(f"创作者 {author} 的数据有效，但可能不足以生成热力图。")
        return

    # 优化热力图可视化
    fig, ax = plt.subplots(figsize=(15, max(5, len(pivot.index) // 2)))
    sns.heatmap(pivot, cmap='YlGnBu', ax=ax, linewidths=.5, annot=True, fmt=".0f", annot_kws={"size": 8})
    ax.set_title(f"{author} 内容发布热力图 (按小时播放量总和)", fontsize=14)
    ax.set_xlabel("小时 (24H)", fontsize=12)
    ax.set_ylabel("日期", fontsize=12)
    plt.xticks(rotation=0)
    plt.yticks(rotation=0)
    st.pyplot(fig)

    st.dataframe(author_df[['title', 'tags_list', 'created_at', 'views', 'engagement']].sort_values('created_at', ascending=False))


# --- 报告自动洞察 ---
def generate_insight(df):
    """基于规则和数据归纳自动生成本期内容洞察"""
    total_videos = len(df)
    top_tag = df.explode('tags_list')['tags_list'].value_counts().idxmax()
    top_author = df['author'].value_counts().idxmax()
    avg_views = df['views'].mean()
    avg_engage = df['engagement'].mean()
    insight = f"本期共分析{total_videos}条视频，最热门标签为【{top_tag}】，头部创作者为【{top_author}】。平均播放量{avg_views:,.0f}，平均互动量{avg_engage:,.0f}。建议关注高潜力标签和头部账号的内容策略，结合趋势预测和标签竞争度分析优化选题和标签组合。"
    return insight


# ==============================================================================
# 5. 主程序入口
# ==============================================================================
def get_data_file_path():
    """
    动态查找数据文件 'bilibili_hot_with_tags.csv'。
    优先在脚本所在目录查找，其次在项目根目录（上一级）查找。
    """
    script_path = Path(__file__).resolve()
    # 定义可能的路径
    current_dir_path = script_path.parent / "bilibili_hot_with_tags.csv"
    parent_dir_path = script_path.parent.parent / "bilibili_hot_with_tags.csv"

    if current_dir_path.exists():
        st.sidebar.info(f"已从当前目录加载数据。")
        return current_dir_path
    elif parent_dir_path.exists():
        st.sidebar.info(f"已从项目根目录加载数据。")
        return parent_dir_path
    else:
        st.sidebar.error("错误：未找到数据文件 'bilibili_hot_with_tags.csv'。")
        st.sidebar.info("请将数据文件放置在与脚本相同的 '专业实习' 目录，或项目的根目录中。")
        return None


def main():
    """主应用控制器。"""
    st.sidebar.title("B站内容决策工具 v3.0")
    
    # 使用健壮的方式获取数据文件路径
    data_path = get_data_file_path()
    if data_path is None:
        # 如果找不到文件，则停止执行
        return

    df = load_data(data_path)
    
    # 增加数据加载校验
    if df is None or df.empty:
        st.error("数据加载失败，请检查数据文件内容或格式。")
        return

    st.sidebar.title("导航菜单")
    page_options = {
        "📈 热点趋势总览": page_overview,
        "📊 趋势预测": page_trend_forecast,
        "🏷️ 标签竞争度": tag_competition_analysis,
        "📅 竞品内容日历": competitor_calendar,
        "🔑 爆款密码挖掘": page_association_rules,
        "🎯 内容策略分析": page_strategy_analysis,
        "🤖 AI模型中心": page_models_center,
        "📄 智能报告生成": page_report_generation
    }
    selection = st.sidebar.radio("选择分析模块", list(page_options.keys()))
    page_func = page_options[selection]
    page_func(df)


if __name__ == "__main__":
    main()